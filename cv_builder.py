"""
cv_builder.py
--------------
Formulaire de création de CV pour l'application "Aide Conseil Emploi".
Mise en page à deux colonnes (bandeau latéral coloré + colonne principale),
avec 3 thèmes de couleur au choix. Pas de photo (décision produit actuelle).

Intégration dans app.py :
    from cv_builder import afficher_generateur_cv
    ...
    with tab_cv:
        afficher_generateur_cv()
"""

import streamlit as st
import re
import os
import requests
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from moteur_recherche import (
    DEPARTEMENTS_VERS_NOM,
    suggerer_postes,
    get_referentiel_appellations,
    _extraire_code_rome,
    resoudre_codes_rome,
    diagnostiquer_romeo,
)
from io import BytesIO


# ---------------------------------------------------------------------------
# Thèmes de couleur
# ---------------------------------------------------------------------------
THEMES = {
    "🔵 Bleu classique": {
        "accent": "2E74B5",       # titres, nom, filets — couleur exacte de la référence
        "bandeau_fond": "EAF1F8",  # fond du bandeau latéral
        "bandeau_texte": "2E74B5",
    },
    "🍷 Bordeaux élégant": {
        "accent": "7B2C3B",
        "bandeau_fond": "F6ECEE",
        "bandeau_texte": "7B2C3B",
    },
    "🟢 Vert forêt": {
        "accent": "2F5233",
        "bandeau_fond": "EAF2EA",
        "bandeau_texte": "2F5233",
    },
}


# ---------------------------------------------------------------------------
# Traduction du CV (FR / EN / ES)
# ---------------------------------------------------------------------------
LANGUES_CV = {
    "🇫🇷 Français": "FR",
    "🇬🇧 English": "EN-GB",
    "🇪🇸 Español": "ES",
}

# Libellés de section fixes — pas besoin d'appel API, ils ne changent jamais.
LIBELLES = {
    "FR": {
        "contact": "Contact", "email": "E-mail", "telephone": "Téléphone", "adresse": "Adresse",
        "langues": "Langues", "competences": "Compétences", "outils": "Outils informatiques",
        "langages": "Langages informatiques", "certifications": "Certifications", "interets": "Centres d'intérêt",
        "experiences": "Expériences professionnelles", "formation": "Formation",
        "presentation": "Présentation", "disponibilite": "Disponibilité",
    },
    "EN-GB": {
        "contact": "Contact", "email": "Email", "telephone": "Phone", "adresse": "Address",
        "langues": "Languages", "competences": "Skills", "outils": "IT Tools",
        "langages": "Programming Languages", "certifications": "Certifications", "interets": "Interests",
        "experiences": "Professional Experience", "formation": "Education",
        "presentation": "Profile", "disponibilite": "Availability",
    },
    "ES": {
        "contact": "Contacto", "email": "Correo electrónico", "telephone": "Teléfono", "adresse": "Dirección",
        "langues": "Idiomas", "competences": "Competencias", "outils": "Herramientas informáticas",
        "langages": "Lenguajes informáticos", "certifications": "Certificaciones", "interets": "Intereses",
        "experiences": "Experiencia profesional", "formation": "Formación",
        "presentation": "Presentación", "disponibilite": "Disponibilidad",
    },
}


def _traduire_lot(textes, langue_cible):
    """
    Traduit une liste de textes en un seul appel DeepL (économise les appels et
    la latence). Dégradation silencieuse vers le texte original (français) si la
    clé API n'est pas configurée ou si l'appel échoue — ne bloque jamais la
    génération du CV.
    """
    if langue_cible == "FR":
        return textes

    cle_api = os.environ.get("DEEPL_API_KEY")
    if not cle_api:
        return textes

    index_non_vides = [i for i, t in enumerate(textes) if t and t.strip()]
    if not index_non_vides:
        return textes

    try:
        url = "https://api-free.deepl.com/v2/translate"
        headers = {"Authorization": f"DeepL-Auth-Key {cle_api}"}
        data = [("text", textes[i]) for i in index_non_vides]
        data += [("target_lang", langue_cible), ("source_lang", "FR")]
        r = requests.post(url, headers=headers, data=data, timeout=15)
        if r.status_code != 200:
            return textes
        traductions = r.json().get("translations", [])
        if len(traductions) != len(index_non_vides):
            return textes
        resultats = list(textes)
        for position, index_original in enumerate(index_non_vides):
            resultats[index_original] = traductions[position]["text"]
        return resultats
    except Exception:
        return textes


# ---------------------------------------------------------------------------
# Drapeaux (facultatif, best-effort — langues reconnues seulement)
# ---------------------------------------------------------------------------
DRAPEAUX_LANGUES = {
    "français": "🇫🇷", "anglais": "🇬🇧", "espagnol": "🇪🇸", "allemand": "🇩🇪",
    "italien": "🇮🇹", "portugais": "🇵🇹", "arabe": "🇸🇦", "chinois": "🇨🇳",
    "mandarin": "🇨🇳", "japonais": "🇯🇵", "russe": "🇷🇺", "néerlandais": "🇳🇱",
    "coréen": "🇰🇷", "turc": "🇹🇷", "polonais": "🇵🇱", "grec": "🇬🇷",
    "hébreu": "🇮🇱", "hindi": "🇮🇳", "suédois": "🇸🇪", "norvégien": "🇳🇴",
    "danois": "🇩🇰", "finnois": "🇫🇮", "roumain": "🇷🇴", "ukrainien": "🇺🇦",
}


def _drapeau_pour_langue(texte_ligne):
    """Retourne un drapeau si le nom de la langue est reconnu, sinon chaîne vide."""
    debut = texte_ligne.strip().lower()
    for nom, drapeau in DRAPEAUX_LANGUES.items():
        if debut.startswith(nom):
            return drapeau + " "
    return ""


# ---------------------------------------------------------------------------
# Échelle automatique (police / espacement) pour tenir sur une page
# ---------------------------------------------------------------------------
def _estimer_volume_contenu(data):
    volume = len(data.get("profil", ""))
    for exp in data.get("experiences", []):
        volume += len(exp.get("description", "")) + 60
    for form in data.get("formations", []):
        volume += 40
    volume += len(data.get("langues", ""))
    volume += len(data.get("competences", ""))
    volume += len(data.get("outils", ""))
    volume += len(data.get("langages_informatiques", ""))
    volume += len(data.get("certifications", ""))
    volume += len(data.get("interets", ""))
    return volume


def _calculer_echelle(volume):
    """Plus le contenu est volumineux, plus on réduit polices/espacements."""
    if volume < 1400:
        return 1.0
    elif volume < 2200:
        return 0.92
    elif volume < 3000:
        return 0.85
    elif volume < 3800:
        return 0.78
    elif volume < 4600:
        return 0.71
    elif volume < 5500:
        return 0.65
    elif volume < 6500:
        return 0.60
    else:
        return 0.55


def _calculer_marge_verticale(echelle):
    """Réduit aussi les marges haut/bas de page pour les contenus très volumineux."""
    if echelle >= 0.85:
        return 1.2
    elif echelle >= 0.65:
        return 0.9
    else:
        return 0.6


def _pt(base, echelle):
    return Pt(round(base * echelle * 2) / 2)


# ---------------------------------------------------------------------------
# Initialisation de l'état
# ---------------------------------------------------------------------------
def _init_cv_state():
    if "cv_experiences" not in st.session_state:
        st.session_state.cv_experiences = []
    if "cv_formations" not in st.session_state:
        st.session_state.cv_formations = []


# ---------------------------------------------------------------------------
# Sections dynamiques (expériences / formations)
# ---------------------------------------------------------------------------
def _section_experiences():
    st.markdown("#### 💼 Expériences professionnelles")

    a_supprimer = None
    for i, exp in enumerate(st.session_state.cv_experiences):
        with st.container(border=True):
            c1, c2 = st.columns(2)
            exp["poste"] = c1.text_input("Poste", value=exp.get("poste", ""), key=f"exp_poste_{i}")
            exp["entreprise"] = c2.text_input("Entreprise", value=exp.get("entreprise", ""), key=f"exp_entreprise_{i}")

            c3, c4, c5, c6 = st.columns(4)
            exp["ville"] = c3.text_input("Ville", value=exp.get("ville", ""), key=f"exp_ville_{i}")
            exp["pays"] = c4.text_input("Pays", value=exp.get("pays", ""), key=f"exp_pays_{i}")
            exp["date_debut"] = c5.text_input("Début (ex: Jan. 2022)", value=exp.get("date_debut", ""), key=f"exp_debut_{i}")
            exp["date_fin"] = c6.text_input("Fin (ex: Déc. 2023 ou En cours)", value=exp.get("date_fin", ""), key=f"exp_fin_{i}")

            exp["description"] = st.text_area(
                "Missions / réalisations (une ligne = une puce)",
                value=exp.get("description", ""),
                key=f"exp_description_{i}",
                height=100,
            )

            if st.button("🗑️ Supprimer cette expérience", key=f"exp_supprimer_{i}"):
                a_supprimer = i

    if a_supprimer is not None:
        st.session_state.cv_experiences.pop(a_supprimer)
        st.rerun()

    if st.button("➕ Ajouter une expérience"):
        st.session_state.cv_experiences.append({})
        st.rerun()


def _section_formations():
    st.markdown("#### 🎓 Formation")

    a_supprimer = None
    for i, form in enumerate(st.session_state.cv_formations):
        with st.container(border=True):
            c1, c2 = st.columns(2)
            form["diplome"] = c1.text_input("Diplôme", value=form.get("diplome", ""), key=f"form_diplome_{i}")
            form["etablissement"] = c2.text_input("Établissement", value=form.get("etablissement", ""), key=f"form_etab_{i}")

            c3, c4, c5 = st.columns(3)
            form["ville"] = c3.text_input("Ville", value=form.get("ville", ""), key=f"form_ville_{i}")
            form["pays"] = c4.text_input("Pays", value=form.get("pays", ""), key=f"form_pays_{i}")
            form["annee"] = c5.text_input("Année (ex: sept 2017 / oct 2018)", value=form.get("annee", ""), key=f"form_annee_{i}")

            if st.button("🗑️ Supprimer cette formation", key=f"form_supprimer_{i}"):
                a_supprimer = i

    if a_supprimer is not None:
        st.session_state.cv_formations.pop(a_supprimer)
        st.rerun()

    if st.button("➕ Ajouter une formation"):
        st.session_state.cv_formations.append({})
        st.rerun()


# ---------------------------------------------------------------------------
# Helpers python-docx bas niveau (ombrage de cellule, bordures de tableau)
# ---------------------------------------------------------------------------
def _ombrer_cellule(cell, couleur_hex):
    """Applique une couleur de fond à une cellule de tableau (non exposé par l'API haut niveau)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), couleur_hex)
    tc_pr.append(shd)


def _supprimer_bordures_tableau(table):
    """Retire toutes les bordures d'un tableau (utilisé comme grille de mise en page invisible)."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for cote in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{cote}")
        elem.set(qn("w:val"), "nil")
        borders.append(elem)
    tbl_pr.append(borders)


def _definir_marges_cellule(cell, gauche=0.15, droite=0.15, haut=0.05, bas=0.05):
    """Définit des marges internes (en cm) pour une cellule."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for cote, valeur in (("left", gauche), ("right", droite), ("top", haut), ("bottom", bas)):
        elem = OxmlElement(f"w:{cote}")
        elem.set(qn("w:w"), str(int(valeur * 567)))  # cm -> twips (1cm ≈ 567 twips)
        elem.set(qn("w:type"), "dxa")
        tc_mar.append(elem)
    tc_pr.append(tc_mar)


def _titre_section(cell_ou_doc, texte, couleur_hex, taille=12, echelle=1.0, espace_avant=12):
    """Ajoute un titre de section stylé (majuscules, gras, coloré)."""
    p = cell_ou_doc.add_paragraph()
    p.paragraph_format.space_before = _pt(espace_avant, echelle)
    p.paragraph_format.space_after = _pt(4, echelle)
    run = p.add_run(texte.upper())
    run.bold = True
    run.font.size = _pt(taille, echelle)
    run.font.color.rgb = RGBColor.from_string(couleur_hex)
    return p


_CARACTERES_PUCE_PARASITES = " -•➤▸●○*>·‣¬▪"


def _nettoyer_ligne(texte):
    """Retire les puces/symboles que l'utilisateur a pu coller depuis un autre CV
    (➤, ▸, -, •...) pour éviter un double affichage avec notre propre puce."""
    return texte.strip(_CARACTERES_PUCE_PARASITES).strip()


def _puce(cell_ou_doc, texte, couleur_puce=None, taille=10, echelle=1.0, caractere="¬"):
    p = cell_ou_doc.add_paragraph()
    # Espacement conforme à la valeur mesurée dans le format de référence (~3pt)
    p.paragraph_format.space_after = _pt(3, echelle)
    run = p.add_run(f"{caractere} {texte}")
    run.font.size = _pt(taille, echelle)
    if couleur_puce:
        run.font.color.rgb = RGBColor.from_string(couleur_puce)
    return p


# ---------------------------------------------------------------------------
# Tri automatique par date (expériences / formations, anti-chronologique)
# ---------------------------------------------------------------------------
_MOIS_FR_NUM = {
    "janvier": 1, "jan": 1, "février": 2, "fevrier": 2, "fév": 2, "fev": 2,
    "mars": 3, "avril": 4, "avr": 4, "mai": 5, "juin": 6, "juillet": 7, "juil": 7,
    "août": 8, "aout": 8, "septembre": 9, "sept": 9, "sep": 9,
    "octobre": 10, "oct": 10, "novembre": 11, "nov": 11,
    "décembre": 12, "decembre": 12, "déc": 12, "dec": 12,
}


def _valeur_tri_date(texte):
    """
    Extrait une valeur triable (année x 12 + mois) à partir d'un texte de date
    libre en français (ex: "Août 2025", "sept 2017 / oct 2018", "En cours").
    "En cours" est traité comme la date la plus récente possible. Retourne 0
    si aucune date n'est reconnue (l'élément descend en fin de liste).
    """
    if not texte:
        return 0
    texte_normalise = texte.lower().strip()
    if any(mot in texte_normalise for mot in ("en cours", "aujourd'hui", "present", "présent")):
        return 999999

    annees = [int(a) for a in re.findall(r"(?:19|20)\d{2}", texte_normalise)]
    if not annees:
        return 0
    annee_retenue = max(annees)  # la plus tardive mentionnée (ex: "sept 2017 / oct 2018" -> 2018)

    mois_retenu = 1
    for nom_mois, num_mois in _MOIS_FR_NUM.items():
        if nom_mois in texte_normalise:
            mois_retenu = num_mois

    return annee_retenue * 12 + mois_retenu


def _trier_par_date(elements, cle_principale, cle_secondaire=None):
    """Trie une liste d'expériences/formations du plus récent au plus ancien."""

    def _cle_tri(element):
        texte = element.get(cle_principale, "")
        if not texte and cle_secondaire:
            texte = element.get(cle_secondaire, "")
        return _valeur_tri_date(texte)

    return sorted(elements, key=_cle_tri, reverse=True)


# ---------------------------------------------------------------------------
# Majuscule automatique (première lettre uniquement, conventions françaises)
# ---------------------------------------------------------------------------
def _majuscule_premiere_lettre(texte):
    texte = (texte or "").strip()
    if not texte:
        return texte
    return texte[0].upper() + texte[1:]


# ---------------------------------------------------------------------------
# Génération du document Word (mise en page 2 colonnes)
# ---------------------------------------------------------------------------
def generer_cv_docx(data, theme_nom="🔵 Bleu classique", photo_bytes=None, afficher_drapeaux=True, langue="FR"):
    theme = THEMES.get(theme_nom, THEMES["🔵 Bleu classique"])
    accent = theme["accent"]
    bandeau_fond = theme["bandeau_fond"]
    bandeau_texte = theme["bandeau_texte"]
    echelle = _calculer_echelle(_estimer_volume_contenu(data))
    libelles = LIBELLES.get(langue, LIBELLES["FR"])

    # --- Traduction groupée des champs texte libre (un seul appel API DeepL) ---
    experiences_brutes = data.get("experiences", [])
    formations_brutes = data.get("formations", [])

    textes_a_traduire = [data.get("profil", ""), data.get("titre_recherche", "")]
    for exp in experiences_brutes:
        textes_a_traduire.append(exp.get("poste", ""))
        textes_a_traduire.append(exp.get("description", ""))
    for form in formations_brutes:
        textes_a_traduire.append(form.get("diplome", ""))

    textes_traduits = _traduire_lot(textes_a_traduire, langue)

    experiences_traduites = []
    formations_traduites = []
    curseur = 2
    for exp in experiences_brutes:
        exp_copie = dict(exp)
        exp_copie["poste"] = textes_traduits[curseur]
        exp_copie["description"] = textes_traduits[curseur + 1]
        curseur += 2
        experiences_traduites.append(exp_copie)
    for form in formations_brutes:
        form_copie = dict(form)
        form_copie["diplome"] = textes_traduits[curseur]
        curseur += 1
        formations_traduites.append(form_copie)

    # On poursuit sur une copie de data avec les champs traduits substitués —
    # noms, dates, villes, pays, e-mail... restent inchangés (non traduits).
    data = dict(data)
    data["profil"] = textes_traduits[0]
    data["titre_recherche"] = textes_traduits[1]
    data["experiences"] = experiences_traduites
    data["formations"] = formations_traduites

    doc = Document()
    # Interligne compact par défaut (évite l'espacement 1.08/1.15 par défaut de Word,
    # qui gonfle inutilement la hauteur de chaque ligne de texte).
    style_normal = doc.styles["Normal"]
    style_normal.paragraph_format.line_spacing = 1.0
    style_normal.paragraph_format.space_after = Pt(0)

    for section in doc.sections:
        section.top_margin = Cm(0.46)
        section.bottom_margin = Cm(0.46)
        section.left_margin = Cm(0.46)
        section.right_margin = Cm(0.46)
        largeur_utile = section.page_width - section.left_margin - section.right_margin

    largeur_bandeau = Inches(2.3)
    largeur_principale = largeur_utile - largeur_bandeau

    # --- Tableau de mise en page 1 ligne x 2 colonnes, bordures invisibles ---
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    _supprimer_bordures_tableau(table)
    table.columns[0].width = largeur_bandeau
    table.columns[1].width = largeur_principale

    cell_bandeau = table.cell(0, 0)
    cell_principale = table.cell(0, 1)
    cell_bandeau.width = largeur_bandeau
    cell_principale.width = largeur_principale
    _ombrer_cellule(cell_bandeau, bandeau_fond)
    marge_cellule_haut = 0.15 if echelle < 0.85 else 0.3
    _definir_marges_cellule(cell_bandeau, gauche=0.35, droite=0.25, haut=marge_cellule_haut, bas=0.15)
    _definir_marges_cellule(cell_principale, gauche=0.35, droite=0.1, haut=marge_cellule_haut, bas=0.15)
    cell_bandeau.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell_principale.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Suppression réelle (pas juste vidage du texte) du paragraphe auto-créé dans
    # chaque cellule — sinon il reste une ligne vide qui pousse tout le contenu
    # vers le bas, y compris le nom en haut de la colonne principale.
    for cellule in (cell_bandeau, cell_principale):
        p_vide = cellule.paragraphs[0]
        p_vide._element.getparent().remove(p_vide._element)

    # =======================================================================
    # BANDEAU LATÉRAL
    # =======================================================================
    # --- Photo (facultative) ---
    if photo_bytes:
        p_photo = cell_bandeau.add_paragraph()
        p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_photo.paragraph_format.space_after = _pt(10, echelle)
        run_photo = p_photo.add_run()
        run_photo.add_picture(BytesIO(photo_bytes), width=Inches(1.6))

    # --- Contact ---
    _titre_section(cell_bandeau, libelles["contact"], bandeau_texte, echelle=echelle, espace_avant=0)
    for icone, label, valeur in [
        ("📧", libelles["email"], data.get("email")),
        ("📱", libelles["telephone"], data.get("telephone")),
        ("🏠", libelles["adresse"], data.get("adresse")),
    ]:
        if valeur:
            p = cell_bandeau.add_paragraph()
            p.paragraph_format.space_after = _pt(6, echelle)
            run_label = p.add_run(f"{icone} {label}\n")
            run_label.bold = True
            run_label.font.size = _pt(9, echelle)
            run_label.font.color.rgb = RGBColor.from_string(bandeau_texte)
            run_val = p.add_run(valeur)
            run_val.font.size = _pt(10, echelle)

    # --- Langues ---
    if data.get("langues"):
        _titre_section(cell_bandeau, libelles["langues"], bandeau_texte, echelle=echelle)
        for ligne in data["langues"].split("\n"):
            ligne = _nettoyer_ligne(ligne)
            if ligne:
                prefixe = _drapeau_pour_langue(ligne) if afficher_drapeaux else ""
                _puce(cell_bandeau, f"{prefixe}{ligne}", echelle=echelle, caractere="▪")

    # --- Compétences ---
    if data.get("competences"):
        _titre_section(cell_bandeau, libelles["competences"], bandeau_texte, echelle=echelle)
        for ligne in data["competences"].split("\n"):
            ligne = _nettoyer_ligne(ligne)
            if ligne:
                _puce(cell_bandeau, ligne, echelle=echelle, caractere="▪")

    # --- Outils informatiques ---
    if data.get("outils"):
        _titre_section(cell_bandeau, libelles["outils"], bandeau_texte, echelle=echelle)
        for ligne in data["outils"].split("\n"):
            ligne = _nettoyer_ligne(ligne)
            if ligne:
                _puce(cell_bandeau, ligne, echelle=echelle, caractere="▪")

    # --- Langages informatiques (facultatif, invisible si vide — profils non-tech) ---
    if data.get("langages_informatiques"):
        _titre_section(cell_bandeau, libelles["langages"], bandeau_texte, echelle=echelle)
        for ligne in data["langages_informatiques"].split("\n"):
            ligne = _nettoyer_ligne(ligne)
            if ligne:
                _puce(cell_bandeau, ligne, echelle=echelle, caractere="▪")

    # --- Certifications (facultatif, invisible si vide) ---
    if data.get("certifications"):
        _titre_section(cell_bandeau, libelles["certifications"], bandeau_texte, echelle=echelle)
        for ligne in data["certifications"].split("\n"):
            ligne = _nettoyer_ligne(ligne)
            if ligne:
                _puce(cell_bandeau, ligne, echelle=echelle, caractere="▪")

    # --- Centres d'intérêt ---
    if data.get("interets"):
        _titre_section(cell_bandeau, libelles["interets"], bandeau_texte, echelle=echelle)
        interets_list = [i.strip() for i in data["interets"].replace("\n", ",").split(",") if i.strip()]
        for interet in interets_list:
            _puce(cell_bandeau, interet, echelle=echelle, caractere="▪")

    # =======================================================================
    # COLONNE PRINCIPALE
    # =======================================================================
    # --- En-tête : nom + titre recherché ---
    p_nom = cell_principale.add_paragraph()
    p_nom.paragraph_format.space_after = Pt(0)
    run_nom = p_nom.add_run(f"{data.get('prenom', '')} {data.get('nom', '')}".strip().upper())
    run_nom.bold = True
    run_nom.font.size = _pt(28, echelle)
    run_nom.font.color.rgb = RGBColor.from_string(accent)

    if data.get("titre_recherche"):
        p_titre = cell_principale.add_paragraph()
        p_titre.paragraph_format.space_after = _pt(8, echelle)
        run_titre = p_titre.add_run(data["titre_recherche"])
        run_titre.italic = True
        run_titre.font.size = _pt(13, echelle)
        run_titre.font.color.rgb = RGBColor.from_string(accent)

    # Filet horizontal sous l'en-tête
    p_filet = cell_principale.add_paragraph()
    p_filet.paragraph_format.space_after = _pt(8, echelle)
    pPr = p_filet._p.get_or_add_pPr()
    bord = OxmlElement("w:pBdr")
    bas = OxmlElement("w:bottom")
    bas.set(qn("w:val"), "single")
    bas.set(qn("w:sz"), "12")
    bas.set(qn("w:space"), "1")
    bas.set(qn("w:color"), accent)
    bord.append(bas)
    pPr.append(bord)

    # --- Présentation (label en gras intégré au paragraphe, pas de titre de section séparé) ---
    if data.get("profil"):
        p = cell_principale.add_paragraph()
        p.paragraph_format.space_after = _pt(4, echelle)
        run_label = p.add_run(f"{libelles['presentation']} : ")
        run_label.bold = True
        run_label.font.size = _pt(10.5, echelle)
        run_texte = p.add_run(data["profil"])
        run_texte.font.size = _pt(10.5, echelle)

    # --- Disponibilité ---
    if data.get("disponibilite"):
        p_dispo = cell_principale.add_paragraph()
        p_dispo.paragraph_format.space_after = _pt(8, echelle)
        run_dispo_label = p_dispo.add_run(f"{libelles['disponibilite']} : ")
        run_dispo_label.bold = True
        run_dispo_label.italic = True
        run_dispo_label.font.size = _pt(10, echelle)
        run_dispo_val = p_dispo.add_run(data["disponibilite"])
        run_dispo_val.italic = True
        run_dispo_val.font.size = _pt(10, echelle)

    # --- Expériences ---
    experiences = [e for e in data.get("experiences", []) if e.get("poste") or e.get("entreprise")]
    experiences = _trier_par_date(experiences, "date_fin", "date_debut")
    if experiences:
        _titre_section(cell_principale, libelles["experiences"], accent, taille=16, echelle=echelle)
        for exp in experiences:
            p = cell_principale.add_paragraph()
            p.paragraph_format.space_before = _pt(6, echelle)
            p.paragraph_format.space_after = Pt(0)
            poste_maj = _majuscule_premiere_lettre(exp.get("poste", ""))
            run = p.add_run(poste_maj)
            run.bold = True
            run.font.size = _pt(11, echelle)

            dates = f"{exp.get('date_debut', '')} - {exp.get('date_fin', '')}".strip(" -")
            meta_parties = [x for x in [exp.get("entreprise", ""), dates] if x]
            meta_texte = " | ".join(meta_parties)
            lieu_pays = " · ".join(x for x in [exp.get("ville", ""), exp.get("pays", "")] if x)
            if lieu_pays:
                meta_texte = f"{meta_texte} · {lieu_pays}" if meta_texte else lieu_pays

            if meta_texte:
                p_meta = cell_principale.add_paragraph()
                p_meta.paragraph_format.space_after = _pt(3, echelle)
                run_meta = p_meta.add_run(meta_texte)
                run_meta.italic = True
                run_meta.font.size = _pt(9.5, echelle)
                run_meta.font.color.rgb = RGBColor.from_string(accent)

            description = exp.get("description", "").strip()
            if description:
                for ligne in description.split("\n"):
                    ligne = _nettoyer_ligne(ligne)
                    if ligne:
                        _puce(cell_principale, ligne, taille=10, echelle=echelle)

    # --- Formation ---
    formations = [f for f in data.get("formations", []) if f.get("diplome") or f.get("etablissement")]
    formations = _trier_par_date(formations, "annee")
    if formations:
        _titre_section(cell_principale, libelles["formation"], accent, taille=16, echelle=echelle)
        for form in formations:
            p = cell_principale.add_paragraph()
            p.paragraph_format.space_before = _pt(4, echelle)
            p.paragraph_format.space_after = Pt(0)
            diplome_maj = _majuscule_premiere_lettre(form.get("diplome", ""))
            run = p.add_run(f"{diplome_maj} — {form.get('etablissement', '')}")
            run.bold = True
            run.font.size = _pt(10.5, echelle)

            meta = " · ".join(
                x for x in [form.get("annee", ""), form.get("ville", ""), form.get("pays", "")] if x
            )
            if meta:
                p_meta = cell_principale.add_paragraph()
                p_meta.paragraph_format.space_after = _pt(2, echelle)
                run_meta = p_meta.add_run(meta)
                run_meta.italic = True
                run_meta.font.size = _pt(9.5, echelle)
                run_meta.font.color.rgb = RGBColor.from_string(accent)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, echelle



# ---------------------------------------------------------------------------
# Suggestions de compétences / outils / langages (basées sur les offres réelles)
# ---------------------------------------------------------------------------
def _ajouter_suggestion(cle_session, valeur):
    """Ajoute une ligne à un textarea (par clé de session_state) si elle n'y est pas déjà."""
    actuel = st.session_state.get(cle_session, "")
    lignes = [l.strip() for l in actuel.split("\n") if l.strip()]
    if valeur not in lignes:
        lignes.append(valeur)
        st.session_state[cle_session] = "\n".join(lignes)


def _ajouter_suggestion(cle_session, valeur):
    """Ajoute une valeur à une liste d'options (session_state) si elle n'y est pas déjà."""
    if cle_session not in st.session_state:
        st.session_state[cle_session] = []
    if valeur not in st.session_state[cle_session]:
        st.session_state[cle_session].append(valeur)


# Listes de base toujours proposées, même sans recherche de profil préalable
# (pour que le champ ne soit jamais "vide" par défaut).
_DEFAUTS_COMPETENCES = [
    "Gestion de projet", "Communication", "Travail d'équipe", "Résolution de problèmes",
    "Organisation", "Leadership", "Analyse", "Négociation", "Gestion du temps", "Esprit critique",
]
_DEFAUTS_OUTILS = ["Excel", "Word", "PowerPoint", "Outlook", "Teams"]
_DEFAUTS_LANGAGES = []  # vide par défaut : pertinent seulement pour les profils tech
_DEFAUTS_CERTIFICATIONS = []  # vide par défaut : très spécifique au poste, pas de base générique pertinente
# Top 10 des langues les plus parlées au monde (nombre total de locuteurs, classement
# usuel type Ethnologue) — liste de départ ; le champ permet aussi d'en ajouter d'autres.
_DEFAUTS_LANGUES = [
    "Anglais", "Mandarin", "Hindi", "Espagnol", "Français",
    "Arabe", "Bengali", "Russe", "Portugais", "Ourdou",
]


def _champ_liste_avec_ajout(titre, cle_base, valeurs_par_defaut, aide=None):
    """
    Affiche une liste à choix multiples (options par défaut + suggestions éventuelles)
    avec possibilité d'ajouter ses propres éléments à la liste. Retourne le résultat
    au format 'une valeur par ligne' (compatible avec generer_cv_docx).
    """
    cle_options = f"{cle_base}_options"
    if cle_options not in st.session_state:
        st.session_state[cle_options] = list(valeurs_par_defaut)

    cle_select = f"{cle_base}_select"
    selection = st.multiselect(titre, options=st.session_state[cle_options], key=cle_select, help=aide)

    col_ajout, col_bouton = st.columns([4, 1])
    nouvel_element = col_ajout.text_input(
        f"Ajouter un élément à « {titre} »",
        key=f"{cle_base}_nouveau",
        label_visibility="collapsed",
        placeholder="Ajouter un élément non listé...",
    )
    if col_bouton.button("➕ Ajouter", key=f"{cle_base}_bouton_ajout"):
        valeur = nouvel_element.strip()
        if valeur:
            _ajouter_suggestion(cle_options, valeur)
            if valeur not in st.session_state[cle_select]:
                st.session_state[cle_select] = st.session_state[cle_select] + [valeur]
            st.rerun()

    return "\n".join(selection)


def _selecteur_poste_recherche(titre_recherche):
    """
    Étiquettes cliquables (multi-sélection) parmi les intitulés ROME proches du texte
    tapé dans "Titre du poste recherché" — un intitulé libre comme "Chef de projet" ne
    correspond souvent à rien de tel quel dans le référentiel France Travail (qui
    distingue "Chef de projet informatique", "Chef de projet BTP"...). Ces étiquettes
    servent à choisir les intitulés réellement interrogés sur la base France Travail,
    pour "Tendance par profil" et les suggestions de compétences/outils/langages
    ci-dessous — pas besoin de ressaisir un poste ailleurs dans l'application.
    """
    appellations = get_referentiel_appellations()
    cle_selection = "cv_postes_recherche_selectionnes"
    cle_terme_precedent = "cv_postes_recherche_terme_precedent"
    cle_auto = "cv_postes_recherche_auto_pour_terme"

    if cle_selection not in st.session_state:
        st.session_state[cle_selection] = []

    suggestions = suggerer_postes(titre_recherche) if titre_recherche.strip() else []

    # Un nouveau terme de recherche efface la sélection précédente : sinon les postes
    # d'une recherche antérieure restent cochés en changeant complètement de sujet.
    terme_precedent = st.session_state.get(cle_terme_precedent, titre_recherche)
    if titre_recherche != terme_precedent and st.session_state[cle_selection]:
        st.session_state[cle_selection] = []
    st.session_state[cle_terme_precedent] = titre_recherche

    # Auto-sélectionne la meilleure suggestion une fois par terme, pour ne pas dépendre
    # d'un clic si l'utilisateur ne remarque pas les étiquettes.
    if suggestions and not st.session_state[cle_selection] and st.session_state.get(cle_auto) != titre_recherche:
        st.session_state[cle_selection].append(suggestions[0])
    st.session_state[cle_auto] = titre_recherche

    if suggestions or st.session_state[cle_selection]:
        tous_les_tags = list(dict.fromkeys(suggestions + st.session_state[cle_selection]))
        colonnes_tags = st.columns(2)
        for i, label in enumerate(tous_les_tags):
            est_selectionne = label in st.session_state[cle_selection]
            texte_bouton = f"✅ {label}" if est_selectionne else label
            col_tag = colonnes_tags[i % 2]
            if col_tag.button(texte_bouton, key=f"cv_poste_tag_{i}_{label}"):
                if est_selectionne:
                    st.session_state[cle_selection].remove(label)
                else:
                    st.session_state[cle_selection].append(label)
                st.rerun()
    elif titre_recherche.strip():
        st.caption("Aucune suggestion trouvée pour ce terme — essaie une autre formulation.")

    postes_choisis = st.session_state[cle_selection]

    # Résolution des codes ROME (même mécanisme que dans l'Espace Candidat).
    codes_par_poste = {}
    departement_pour_resolution = st.session_state.get("cv_departement") or "13"
    for label in postes_choisis:
        item_poste = next((a for a in appellations if a.get("libelle", "").strip() == label), None)
        code = _extraire_code_rome(item_poste) if item_poste else None
        if not code:
            df_resolu = resoudre_codes_rome(mots_cles=label, departement=departement_pour_resolution)
            code = df_resolu.iloc[0]["code_rome"] if not df_resolu.empty else None
        codes_par_poste[label] = code

    st.session_state["cv_postes_recherche"] = postes_choisis
    st.session_state["cv_codes_par_poste"] = codes_par_poste


def _section_suggestions_competences(fonction_analyse_competences):
    """
    Alimente automatiquement les listes de compétences/outils/langages suggérées à
    partir des offres correspondant aux postes sélectionnés juste au-dessus (étiquettes
    ROME) — se déclenche seul, pas besoin de visiter un autre onglet ni de cliquer.
    """
    if not fonction_analyse_competences:
        return

    postes_choisis = st.session_state.get("cv_postes_recherche", [])
    if not postes_choisis:
        st.info(
            "👉 Renseigne un poste ci-dessus et choisis au moins une suggestion pour enrichir "
            "automatiquement ces listes avec les compétences réellement demandées sur ce métier "
            "(sinon, une liste générique de base reste disponible ci-dessous)."
        )
        return

    departement_cv = st.session_state.get("cv_departement") or "13"
    cle_signature = "cv_suggestions_signature"
    signature_actuelle = (tuple(postes_choisis), departement_cv)

    if st.session_state.get(cle_signature) != signature_actuelle:
        with st.spinner("Analyse des offres en cours..."):
            mots_cles_larges = " ".join(postes_choisis)
            df_comp, df_outils, df_langages, df_certifs, nb_total = fonction_analyse_competences(
                "TOUS", departement_cv, mots_cles=mots_cles_larges,
            )
        for cle_options, df in [
            ("cv_competences_options", df_comp),
            ("cv_outils_options", df_outils),
            ("cv_langages_options", df_langages),
            ("cv_certifications_options", df_certifs),
        ]:
            if cle_options not in st.session_state:
                st.session_state[cle_options] = []
            for lib in df["libelle"]:
                if lib not in st.session_state[cle_options]:
                    st.session_state[cle_options].append(lib)
        st.session_state["cv_suggestions_apercu"] = (df_comp, df_outils, df_langages, df_certifs, nb_total)
        st.session_state[cle_signature] = signature_actuelle

    if "cv_suggestions_apercu" in st.session_state:
        df_comp, df_outils, df_langages, df_certifs, nb_total = st.session_state["cv_suggestions_apercu"]
        if nb_total < 10:
            st.caption(f"⚠️ Échantillon réduit ({nb_total} offre(s)) — indicatif seulement.")
        else:
            st.caption(f"✅ Listes enrichies automatiquement à partir de {nb_total} offre(s) trouvée(s).")
        for titre_apercu, df in [
            ("Compétences les + demandées", df_comp),
            ("Outils les + demandés", df_outils),
            ("Langages les + demandés", df_langages),
            ("Certifications les + demandées", df_certifs),
        ]:
            if not df.empty:
                apercu = ", ".join(f"{r.libelle} ({r.pourcentage}%)" for _, r in df.head(6).iterrows())
                st.caption(f"💡 **{titre_apercu}** : {apercu}")


# ---------------------------------------------------------------------------
# Interface Streamlit
# ---------------------------------------------------------------------------
def afficher_generateur_cv(fonction_analyse_competences=None):
    _init_cv_state()

    st.header("🧾 Créez votre CV")
    st.caption("Créez votre CV professionnel, prêt à l'emploi, au format Word.")
    st.markdown(
        "**Comment ça marche ici :** renseignez vos informations ci-dessous (coordonnées, "
        "expériences, formations, compétences...), choisissez un thème de couleur, puis générez "
        "votre CV en un clic."
    )
    st.markdown(
        "**Le parcours complet de l'application :**\n"
        "1. 🧾 **Créer mon CV** *(vous êtes ici)* — construisez votre CV et définissez le poste "
        "que vous visez.\n"
        "2. 🎯 **Tendance par profil** — se lance automatiquement dès que votre poste est "
        "renseigné : tension du marché, villes qui recrutent, top recruteurs à démarcher.\n"
        "3. 🧩 **KPIs avancés** — pour aller plus loin : évolution du marché, salaires, types "
        "de contrat."
    )

    theme_choisi = st.radio(
        "🎨 Thème de couleur",
        list(THEMES.keys()),
        horizontal=True,
        key="cv_theme",
    )

    langue_choisie_label = st.radio(
        "🌍 Langue du CV",
        list(LANGUES_CV.keys()),
        horizontal=True,
        key="cv_langue",
    )
    langue_choisie = LANGUES_CV[langue_choisie_label]
    if langue_choisie != "FR" and not os.environ.get("DEEPL_API_KEY"):
        st.warning(
            "⚠️ La traduction automatique n'est pas configurée pour l'instant (clé DeepL "
            "manquante) — le CV sera généré en français malgré la langue choisie."
        )

    photo_uploadee = st.file_uploader(
        "📷 Photo (facultatif, format carré recommandé)", type=["png", "jpg", "jpeg"], key="cv_photo"
    )
    if photo_uploadee:
        col_apercu, _ = st.columns([1, 4])
        col_apercu.image(photo_uploadee, width=100)

    afficher_drapeaux = st.checkbox("🏳️ Afficher un drapeau à côté des langues reconnues", value=True, key="cv_drapeaux")

    with st.expander("👤 Informations générales", expanded=True):
        c1, c2 = st.columns(2)
        prenom = c1.text_input("Prénom", key="cv_prenom")
        nom = c2.text_input("Nom", key="cv_nom")

        titre_recherche = st.text_input(
            "Titre du poste recherché (ex: PMO Finance)",
            key="cv_titre",
            help="C'est ce titre qui apparaîtra sur ton CV, sous ton nom — peut être personnalisé librement.",
        )
        st.caption(
            "💡 Privilégie un intitulé générique (ex: « Consultant » plutôt que « Consultant PMO "
            "Finance senior confirmé »). Ci-dessous, choisis un ou plusieurs intitulés officiels "
            "France Travail (ROME) proches — ce sont eux qui alimentent l'analyse automatique de "
            "l'onglet **🎯 Tendance par profil** et les suggestions de compétences plus bas."
        )
        _selecteur_poste_recherche(titre_recherche)

        with st.expander("🔧 Diagnostic technique ROMEO 2 (temporaire)"):
            st.caption(
                "Outil de mise au point — teste chaque combinaison endpoint/champ candidate "
                "et affiche la vraie réponse de l'API, pour identifier la bonne configuration."
            )
            if st.button("Lancer le diagnostic", key="btn_diagnostic_romeo"):
                with st.spinner("Test des combinaisons ROMEO en cours..."):
                    resultats_diag = diagnostiquer_romeo(titre_recherche.strip() or "chef de projet")
                st.json(resultats_diag)

        c3, c4 = st.columns(2)
        email = c3.text_input("Email", key="cv_email")
        telephone = c4.text_input("Téléphone", key="cv_telephone")
        adresse = st.text_input(
            "Adresse", key="cv_adresse", placeholder="ex: 6 Calle Cronista Veravens, 3012 Alicante, España"
        )

        options_departement_cv = ["Non renseigné"] + sorted(
            f"{code} - {nom}" for code, nom in DEPARTEMENTS_VERS_NOM.items()
        )
        departement_choisi_cv = st.selectbox(
            "Département de résidence",
            options=options_departement_cv,
            key="cv_departement_label",
            help=(
                "N'apparaît pas sur le CV — préremplit automatiquement le département dans "
                "l'onglet 🎯 Tendance par profil."
            ),
        )
        if departement_choisi_cv != "Non renseigné":
            st.session_state["cv_departement"] = departement_choisi_cv.split(" - ")[0]
        else:
            st.session_state.pop("cv_departement", None)

        profil = st.text_area(
            "Profil / accroche (2-3 phrases qui résument votre parcours et votre projet)",
            key="cv_profil",
            height=100,
        )
        disponibilite = st.text_input(
            "Disponibilité", key="cv_disponibilite", placeholder="ex: immédiate, sous 1 mois..."
        )

    with st.expander("💼 Expériences professionnelles"):
        _section_experiences()

    with st.expander("🎓 Formation"):
        _section_formations()

    with st.expander("🌍 Langues"):
        langues = _champ_liste_avec_ajout(
            "Sélectionne ou ajoute tes langues (précise le niveau via « Ajouter », ex: « Anglais - Courant »)",
            "cv_langues_choix", _DEFAUTS_LANGUES,
        )

    with st.expander("💡 Compétences, outils, langages & certifications"):
        st.caption("ℹ️ Ces éléments apparaîtront sur votre CV, dans le bandeau latéral.")
        _section_suggestions_competences(fonction_analyse_competences)

        st.markdown("###### 🧠 Compétences")
        competences = _champ_liste_avec_ajout(
            "Sélectionne ou ajoute tes compétences", "cv_competences", _DEFAUTS_COMPETENCES
        )

        st.markdown("###### 🛠️ Outils informatiques")
        outils = _champ_liste_avec_ajout(
            "Sélectionne ou ajoute tes outils", "cv_outils", _DEFAUTS_OUTILS
        )

        st.markdown("###### 💻 Langages informatiques")
        st.caption("Facultatif — pertinent surtout pour les profils tech/data.")
        langages_informatiques = _champ_liste_avec_ajout(
            "Sélectionne ou ajoute tes langages", "cv_langages", _DEFAUTS_LANGAGES
        )

        st.markdown("###### 🎓 Certifications")
        st.caption(
            "Facultatif — ex: PMP, Scrum Master, CISSP, AWS Certified, CACES, permis... Les "
            "suggestions ci-dessus (si disponibles) sont repérées par mot-clé dans les offres "
            "réelles, pas un champ officiel dédié côté France Travail."
        )
        certifications = _champ_liste_avec_ajout(
            "Sélectionne ou ajoute tes certifications", "cv_certifications", _DEFAUTS_CERTIFICATIONS
        )

    with st.expander("🎯 Centres d'intérêt"):
        interets = st.text_area(
            "Séparés par une virgule ou une ligne (ex: Kayak, Dessin, Voyages)",
            key="cv_interets",
            height=60,
        )

    st.divider()

    if st.button("📄 Générer mon CV", type="primary"):
        if not nom or not prenom:
            st.error("Merci de renseigner au minimum votre nom et prénom.")
        else:
            data = {
                "nom": nom,
                "prenom": prenom,
                "titre_recherche": titre_recherche,
                "email": email,
                "telephone": telephone,
                "adresse": adresse,
                "profil": profil,
                "disponibilite": disponibilite,
                "experiences": st.session_state.cv_experiences,
                "formations": st.session_state.cv_formations,
                "langues": langues,
                "competences": competences,
                "langages_informatiques": langages_informatiques,
                "outils": outils,
                "certifications": certifications,
                "interets": interets,
            }
            photo_bytes = photo_uploadee.getvalue() if photo_uploadee else None
            message_attente = (
                "Traduction et génération en cours..." if langue_choisie != "FR" else "Génération en cours..."
            )
            with st.spinner(message_attente):
                buffer, echelle = generer_cv_docx(
                    data,
                    theme_nom=theme_choisi,
                    photo_bytes=photo_bytes,
                    afficher_drapeaux=afficher_drapeaux,
                    langue=langue_choisie,
                )
            st.success("Votre CV est prêt !")
            if echelle < 0.85:
                st.warning(
                    "⚠️ Contenu assez volumineux : la police et les espacements ont été "
                    "automatiquement réduits pour essayer de tenir sur une page. Si le rendu "
                    "final dépasse quand même une page, pense à raccourcir certaines descriptions "
                    "d'expérience."
                )
            st.download_button(
                label="⬇️ Télécharger mon CV (.docx)",
                data=buffer,
                file_name=f"CV_{prenom}_{nom}{'' if langue_choisie == 'FR' else '_' + langue_choisie}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
